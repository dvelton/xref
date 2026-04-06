#!/usr/bin/env python3
"""
Xref - Cross-reference and defined term resolver for legal documents.

Parses legal documents, identifies cross-references and defined terms,
and produces an interactive HTML viewer where every reference is hoverable/expandable.

Usage:
    python3 xref.py extract-structure --source <file>
    python3 xref.py find-references --source <file>
    python3 xref.py resolve-references --structure <json> --references <json>
    python3 xref.py fetch-external --citations <json>
    python3 xref.py build --source <file> --output <file> --resolved <json> --externals <json>
    python3 xref.py setup-check
"""

import argparse
import hashlib
import json
import os
import re
import sys
from collections import defaultdict
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import fitz
except ImportError:
    fitz = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import requests
except ImportError:
    requests = None

try:
    from jinja2 import Template
except ImportError:
    Template = None

# Import statute fetchers
try:
    from statutes import (
        fetch_us_code, fetch_cfr, fetch_uk_legislation, fetch_eu_regulation,
        get_cache_key, get_cached, set_cached, FetchError
    )
except ImportError:
    fetch_us_code = None


def error(msg: str, exit_code: int = 1):
    """Print error message and exit."""
    print(f"Error: {msg}", file=sys.stderr)
    sys.exit(exit_code)


def check_core_deps():
    """Check for missing core dependencies."""
    missing = []
    if DocxDocument is None:
        missing.append("python-docx")
    if fitz is None:
        missing.append("pymupdf")
    if BeautifulSoup is None:
        missing.append("beautifulsoup4")
    if requests is None:
        missing.append("requests")
    if Template is None:
        missing.append("jinja2")
    
    if missing:
        error(f"Missing dependencies: {', '.join(missing)}\n" +
              "Run: bash setup.sh or pip install " + " ".join(missing))


def compute_file_hash(filepath: str) -> str:
    """Compute SHA-256 hash of a file."""
    h = hashlib.sha256()
    with open(filepath, 'rb') as f:
        while chunk := f.read(8192):
            h.update(chunk)
    return h.hexdigest()


# ============================================================================
# EXTRACT STRUCTURE
# ============================================================================

def extract_structure_docx(filepath: str) -> Dict[str, Any]:
    """Extract section structure and defined terms from a Word document."""
    if DocxDocument is None:
        error("python-docx is required for Word document parsing")
    
    doc = DocxDocument(filepath)
    sections = []
    defined_terms = []
    numbering_scheme = "decimal"
    
    # Track section hierarchy
    section_stack = []
    current_section_id = None
    
    # Extract sections using heading styles
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        
        if not text:
            continue
        
        # Detect section headings
        level = None
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1
        
        # Also check for numbered sections in text
        section_match = re.match(r'^((?:\d+\.)*\d+)\s+(.+)$', text)
        if section_match:
            section_num = section_match.group(1)
            title = section_match.group(2)
            
            # Determine level from numbering depth
            if level is None:
                level = section_num.count('.') + 1
            
            section = {
                "id": section_num,
                "number": section_num,
                "title": title,
                "page": None,
                "level": level,
                "text": text,
                "children": []
            }
            
            # Build hierarchy
            while section_stack and section_stack[-1]["level"] >= level:
                section_stack.pop()
            
            if section_stack:
                section_stack[-1]["children"].append(section)
            else:
                sections.append(section)
            
            section_stack.append(section)
            current_section_id = section_num
    
    # Extract defined terms using common patterns
    full_text = "\n".join([p.text for p in doc.paragraphs])
    defined_terms = extract_defined_terms(full_text, sections)
    
    return {
        "sections": sections,
        "defined_terms": defined_terms,
        "numbering_scheme": numbering_scheme,
        "total_pages": len(doc.sections) if hasattr(doc, 'sections') else None,
        "parser_confidence": {
            "overall": "high",
            "uncertain_regions": []
        }
    }


def extract_structure_pdf(filepath: str) -> Dict[str, Any]:
    """Extract section structure and defined terms from a PDF."""
    if fitz is None:
        error("pymupdf is required for PDF parsing")
    
    doc = fitz.open(filepath)
    sections = []
    defined_terms = []
    full_text = ""
    
    # Extract text with font information to identify headings
    section_candidates = []
    
    for page_num, page in enumerate(doc, 1):
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if "lines" not in block:
                continue
            
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"].strip()
                    font_size = span["size"]
                    font_flags = span["flags"]
                    
                    full_text += text + " "
                    
                    # Heuristic: larger font or bold = likely heading
                    is_bold = font_flags & 2**4
                    
                    # Check for section numbering patterns
                    section_match = re.match(r'^((?:\d+\.)*\d+)\s+(.+)$', text)
                    if section_match and (font_size > 11 or is_bold):
                        section_num = section_match.group(1)
                        title = section_match.group(2)
                        level = section_num.count('.') + 1
                        
                        section_candidates.append({
                            "id": section_num,
                            "number": section_num,
                            "title": title,
                            "page": page_num,
                            "level": level,
                            "text": text,
                            "font_size": font_size,
                            "children": []
                        })
    
    # Build section hierarchy
    section_stack = []
    for section in section_candidates:
        level = section["level"]
        
        while section_stack and section_stack[-1]["level"] >= level:
            section_stack.pop()
        
        if section_stack:
            section_stack[-1]["children"].append(section)
        else:
            sections.append(section)
        
        section_stack.append(section)
    
    # Extract defined terms
    defined_terms = extract_defined_terms(full_text, sections)
    
    return {
        "sections": sections,
        "defined_terms": defined_terms,
        "numbering_scheme": "decimal",
        "total_pages": len(doc),
        "parser_confidence": {
            "overall": "medium",
            "uncertain_regions": [
                {
                    "description": "PDF parsing uses font-size heuristics. Results may be less reliable than Word documents."
                }
            ]
        }
    }


def extract_defined_terms(text: str, sections: List[Dict]) -> List[Dict[str, Any]]:
    """
    Extract defined terms from text using 5-step priority algorithm.
    
    1. Definitions section patterns
    2. Inline quotation + bold formatting
    3. Parenthetical definitions
    4. Section-scoped definitions
    5. Build index and detect undefined terms
    """
    terms = []
    term_index = {}
    
    # Pattern 1: "Term" means/shall mean...
    for match in re.finditer(r'"([^"]+)"\s+(?:means|shall mean|has the meaning)\s+([^.]+\.)', text, re.IGNORECASE):
        term = match.group(1)
        definition = match.group(2).strip()
        
        if term not in term_index:
            terms.append({
                "term": term,
                "definition": definition,
                "defined_in": None,
                "page": None,
                "scope": "global"
            })
            term_index[term] = True
    
    # Pattern 2: Parenthetical inline definitions (the "Term")
    for match in re.finditer(r'\((?:the\s+)?"([A-Z][^"()]+)"\)', text):
        term = match.group(1)
        
        if term not in term_index:
            # Extract context as definition
            start = max(0, match.start() - 100)
            context = text[start:match.start()].strip()
            
            terms.append({
                "term": term,
                "definition": context[-50:] if context else "See document context",
                "defined_in": None,
                "page": None,
                "scope": "global"
            })
            term_index[term] = True
    
    # Pattern 3: Section-scoped definitions
    for match in re.finditer(r'For purposes of this (?:Section|Article|Clause)\s+(\S+),\s+"([^"]+)"\s+means\s+([^.]+\.)', text, re.IGNORECASE):
        section_ref = match.group(1)
        term = match.group(2)
        definition = match.group(3).strip()
        
        terms.append({
            "term": term,
            "definition": definition,
            "defined_in": section_ref,
            "page": None,
            "scope": f"section:{section_ref}"
        })
        term_index[term] = True
    
    return terms


# ============================================================================
# FIND REFERENCES
# ============================================================================

def find_references(text: str, sections: List[Dict]) -> Dict[str, Any]:
    """Find all cross-references, defined term usages, and external citations."""
    internal_refs = []
    defined_term_usages = []
    external_citations = []
    unresolved = []
    
    ref_id = 0
    
    # Internal cross-references patterns
    patterns = [
        (r'\bSection\s+((?:\d+\.)*\d+(?:\([a-z]+\))?)', 'section'),
        (r'\bArticle\s+((?:[IVX]+|(?:\d+\.)*\d+))', 'article'),
        (r'\bClause\s+((?:\d+\.)*\d+)', 'clause'),
        (r'\b(Schedule|Exhibit|Annex)\s+([A-Z]|\d+)', 'schedule'),
    ]
    
    for pattern, ref_type in patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            ref_id += 1
            target = match.group(1) if ref_type != 'schedule' else match.group(2)
            
            internal_refs.append({
                "id": f"ref-{ref_id}",
                "text": match.group(0),
                "target_section": target,
                "found_in_section": None,
                "page": None,
                "character_offset": match.start(),
                "context": text[max(0, match.start()-50):match.end()+50]
            })
    
    # External citations
    citation_patterns = [
        (r'(\d+)\s+U\.?S\.?C\.?\s+(?:Section\s+|§\s*)?(\d+[a-z]?)', 'us_code'),
        (r'(\d+)\s+C\.?F\.?R\.?\s+(?:Section\s+)?(\d+)(?:\.(\d+))?', 'cfr'),
        (r'(GDPR)\s+Article\s+(\d+)', 'eu_regulation'),
        (r'([A-Z][\w\s]+Act)\s+(\d{4})(?:,?\s+(?:s\.|section)\s*(\d+))?', 'uk_legislation'),
    ]
    
    for pattern, cit_type in citation_patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            ref_id += 1
            
            if cit_type == 'us_code':
                external_citations.append({
                    "id": f"ext-{ref_id}",
                    "text": match.group(0),
                    "citation_type": cit_type,
                    "title": match.group(1),
                    "section": match.group(2),
                    "found_in_section": None,
                    "page": None,
                    "character_offset": match.start()
                })
            elif cit_type == 'cfr':
                external_citations.append({
                    "id": f"ext-{ref_id}",
                    "text": match.group(0),
                    "citation_type": cit_type,
                    "title": match.group(1),
                    "part": match.group(2),
                    "section": match.group(3),
                    "found_in_section": None,
                    "page": None
                })
            elif cit_type == 'eu_regulation':
                external_citations.append({
                    "id": f"ext-{ref_id}",
                    "text": match.group(0),
                    "citation_type": cit_type,
                    "regulation": match.group(1),
                    "article": match.group(2),
                    "found_in_section": None,
                    "page": None
                })
    
    return {
        "internal_references": internal_refs,
        "defined_term_usages": defined_term_usages,
        "external_citations": external_citations,
        "unresolved": unresolved
    }


# ============================================================================
# RESOLVE REFERENCES
# ============================================================================

def resolve_references(structure: Dict, references: Dict) -> Dict[str, Any]:
    """Resolve all references to their targets."""
    sections_by_id = {}
    
    def index_sections(secs, parent_id=None):
        for sec in secs:
            sections_by_id[sec["id"]] = sec
            sections_by_id[sec["number"]] = sec
            if sec.get("children"):
                index_sections(sec["children"], sec["id"])
    
    index_sections(structure["sections"])
    
    resolved = {
        "internal_references": [],
        "defined_terms": list(structure.get("defined_terms", [])),
        "external_citations": references.get("external_citations", []),
        "unresolved": [],
        "circular_chains": []
    }
    
    # Resolve internal references
    for ref in references.get("internal_references", []):
        target = ref["target_section"]
        
        # Try exact match, then case-insensitive
        if target in sections_by_id:
            resolved_ref = {**ref, "resolved": True, "target": sections_by_id[target]}
            resolved["internal_references"].append(resolved_ref)
        else:
            # Try case-insensitive
            found = False
            for sec_id, sec in sections_by_id.items():
                if sec_id.lower() == target.lower():
                    resolved_ref = {**ref, "resolved": True, "target": sec}
                    resolved["internal_references"].append(resolved_ref)
                    found = True
                    break
            
            if not found:
                resolved["unresolved"].append({
                    **ref,
                    "unresolved_type": "broken_reference",
                    "reason": f"No section {target} found in document"
                })
    
    # Resolve circular definitions: replace "has the meaning set forth in
    # Section X" with the actual text from that section
    forward_ref = re.compile(
        r'has the meaning (?:set forth|given|assigned|specified|defined) in '
        r'(?:Section|Clause|Article)\s+((?:\d+\.)*\d+(?:\([a-z]+\))?)',
        re.IGNORECASE
    )
    for term in resolved["defined_terms"]:
        defn = term.get("definition", "")
        m = forward_ref.search(defn)
        if m:
            target_id = m.group(1)
            target_sec = sections_by_id.get(target_id)
            if target_sec and target_sec.get("text"):
                term["definition"] = target_sec["text"]
                term["resolved_from"] = f"Section {target_id}"

    # Expand inline term references so each definition is self-contained.
    # When a definition mentions another defined term, append that term's
    # meaning in parentheses so the reader never has to cross-reference.
    term_defs = {}
    for term in resolved["defined_terms"]:
        term_defs[term["term"]] = term.get("definition", "")

    # Sort by length descending so longer terms match before shorter ones
    # (e.g., "AI Model API Early Access" before "AI Model API")
    terms_longest_first = sorted(term_defs.keys(), key=len, reverse=True)

    def expand_definition(defn: str, visited: set) -> str:
        """Recursively expand defined terms found inside a definition."""
        for t in terms_longest_first:
            if t in visited:
                continue
            pattern = re.compile(r'\b' + re.escape(t) + r'\b')
            m = pattern.search(defn)
            if not m:
                continue
            # Skip if this match is inside a prior "(i.e., ...)" expansion
            prefix = defn[:m.start()]
            if prefix.count("(i.e.,") > prefix.count(")"):
                continue
            # Skip if the term is part of a larger capitalized phrase
            # (e.g., "Agreement" inside "Non-Disclosure Agreement")
            before = defn[max(0, m.start()-1):m.start()]
            if before and before[-1].isalpha():
                continue
            # Check if preceded by a capitalized word that suggests a compound
            # name, like "Non-Disclosure Agreement" or "Early Access Agreement"
            preceding_ctx = defn[max(0, m.start()-30):m.start()]
            if preceding_ctx and re.search(r'[A-Z]\w+[\s-]+$', preceding_ctx):
                # Looks like "SomeWord Agreement" -- likely a compound name
                # Only expand if 'this' or 'the' immediately precedes
                if not re.search(r'\b(?:this|the)\s+$', preceding_ctx, re.IGNORECASE):
                    continue
            inner = term_defs.get(t, "")
            if not inner or inner == defn:
                continue
            expanded_inner = expand_definition(inner, visited | {t})
            if len(expanded_inner) > 200:
                expanded_inner = expanded_inner[:197] + "..."
            replacement = m.group(0) + " (i.e., " + expanded_inner + ")"
            defn = defn[:m.start()] + replacement + defn[m.end():]
        return defn

    for term in resolved["defined_terms"]:
        original = term.get("definition", "")
        expanded = expand_definition(original, {term["term"]})
        if expanded != original:
            term["definition_expanded"] = expanded
    
    return resolved


# ============================================================================
# FETCH EXTERNAL
# ============================================================================

def fetch_external_citations(citations_json: str) -> List[Dict[str, Any]]:
    """Fetch external statutory references."""
    if fetch_us_code is None:
        error("Statute fetcher modules not available")
    
    citations = json.loads(citations_json)
    results = []
    
    for citation in citations:
        cit_type = citation.get("citation_type")
        
        # Check cache first
        cache_key = get_cache_key(cit_type, **{k: v for k, v in citation.items() if k != "citation_type"})
        cached = get_cached(cache_key)
        
        if cached:
            results.append({**cached, "cached": True})
            continue
        
        try:
            if cit_type == "us_code":
                result = fetch_us_code(citation["title"], citation["section"])
            elif cit_type == "cfr":
                result = fetch_cfr(citation["title"], citation["part"], citation.get("section"))
            elif cit_type == "uk_legislation":
                result = fetch_uk_legislation(citation["act_name"], citation["year"], citation.get("section"))
            elif cit_type == "eu_regulation":
                result = fetch_eu_regulation(citation["regulation"], citation["article"])
            else:
                result = {
                    **citation,
                    "fetched": False,
                    "error": "Unsupported citation type"
                }
            
            # Cache successful fetches
            if result.get("fetched"):
                set_cached(cache_key, result)
            
            results.append(result)
            
        except FetchError as e:
            results.append({
                **citation,
                "fetched": False,
                "error": str(e)
            })
    
    return results


# ============================================================================
# BUILD HTML
# ============================================================================

def build_html(source: str, output: str, resolved_json: str, externals_json: str = None,
               title: str = None, compact: bool = False):
    """Build the interactive HTML viewer."""
    if Template is None:
        error("jinja2 is required for HTML generation")
    
    # Load resolved graph
    with open(resolved_json, 'r') as f:
        resolved = json.load(f)
    
    # Load externals if provided
    externals = []
    if externals_json:
        with open(externals_json, 'r') as f:
            externals = json.load(f)
    
    # Read source document text
    source_text = ""
    if source.endswith('.docx'):
        doc = DocxDocument(source)
        source_text = "\n\n".join([p.text for p in doc.paragraphs])
    elif source.endswith('.pdf'):
        doc = fitz.open(source)
        for page in doc:
            source_text += page.get_text()
    
    # Compute file hash for provenance
    file_hash = compute_file_hash(source)
    
    # Load template
    template_path = os.path.join(os.path.dirname(__file__), "templates", "viewer.html")
    with open(template_path, 'r') as f:
        template = Template(f.read())
    
    # Load CSS and JS
    css_path = os.path.join(os.path.dirname(__file__), "templates", "styles.css")
    js_path = os.path.join(os.path.dirname(__file__), "templates", "viewer.js")
    
    with open(css_path, 'r') as f:
        css = f.read()
    with open(js_path, 'r') as f:
        js = f.read()
    
    # Render HTML
    html = template.render(
        title=title or os.path.basename(source),
        source_file=source,
        file_hash=file_hash,
        document_text=source_text,
        sections=resolved.get("sections", []),
        internal_refs=resolved.get("internal_references", []),
        defined_terms=resolved.get("defined_terms", []),
        external_citations=externals if not compact else [],
        unresolved=resolved.get("unresolved", []),
        css=css,
        js=js,
        compact=compact
    )
    
    # Write output
    with open(output, 'w', encoding='utf-8') as f:
        f.write(html)
    
    print(f"Generated: {output}")


# ============================================================================
# SETUP CHECK
# ============================================================================

def setup_check():
    """Check which dependencies are available."""
    print("Xref setup check:\n")
    
    deps = [
        ("python-docx", DocxDocument is not None, "Word document support"),
        ("pymupdf", fitz is not None, "PDF support"),
        ("beautifulsoup4", BeautifulSoup is not None, "HTML parsing"),
        ("requests", requests is not None, "External citation fetching"),
        ("jinja2", Template is not None, "HTML generation"),
    ]
    
    all_ok = True
    for name, available, description in deps:
        status = "✓" if available else "✗"
        print(f"  {status} {name:20s} - {description}")
        if not available:
            all_ok = False
    
    print()
    
    try:
        import playwright
        print("  ✓ playwright           - Web URL support (optional)")
    except ImportError:
        print("  ○ playwright           - Web URL support (optional, not installed)")
    
    print()
    
    if all_ok:
        print("✓ All core dependencies installed")
        return 0
    else:
        print("✗ Missing dependencies. Run: bash setup.sh")
        return 1


# ============================================================================
# CLI
# ============================================================================

def main():
    parser = argparse.ArgumentParser(description="Xref - Legal document cross-reference resolver")
    subparsers = parser.add_subparsers(dest='command', help='Command to run')
    
    # extract-structure
    extract_parser = subparsers.add_parser('extract-structure', help='Extract document structure')
    extract_parser.add_argument('--source', required=True, help='Source document path')
    extract_parser.add_argument('--ai-assist', action='store_true', help='Enable AI-assisted parsing')
    
    # find-references
    find_parser = subparsers.add_parser('find-references', help='Find all references')
    find_parser.add_argument('--source', required=True, help='Source document path')
    
    # resolve-references
    resolve_parser = subparsers.add_parser('resolve-references', help='Resolve references')
    resolve_parser.add_argument('--structure', required=True, help='Structure JSON file')
    resolve_parser.add_argument('--references', required=True, help='References JSON file')
    
    # fetch-external
    fetch_parser = subparsers.add_parser('fetch-external', help='Fetch external citations')
    fetch_parser.add_argument('--citations', required=True, help='Citations JSON string')
    
    # build
    build_parser = subparsers.add_parser('build', help='Build interactive HTML')
    build_parser.add_argument('--source', required=True, help='Source document path')
    build_parser.add_argument('--output', required=True, help='Output HTML path')
    build_parser.add_argument('--resolved', required=True, help='Resolved graph JSON')
    build_parser.add_argument('--externals', help='External citations JSON')
    build_parser.add_argument('--title', help='Document title')
    build_parser.add_argument('--compact', action='store_true', help='Compact mode (no external text)')
    
    # setup-check
    subparsers.add_parser('setup-check', help='Check dependencies')
    
    args = parser.parse_args()
    
    if not args.command:
        parser.print_help()
        return 1
    
    if args.command == 'setup-check':
        return setup_check()
    
    check_core_deps()
    
    if args.command == 'extract-structure':
        source = args.source
        
        if source.endswith('.docx'):
            result = extract_structure_docx(source)
        elif source.endswith('.pdf'):
            result = extract_structure_pdf(source)
        else:
            error("Unsupported file format. Use .docx or .pdf")
        
        print(json.dumps(result, indent=2))
    
    elif args.command == 'find-references':
        source = args.source
        
        # Extract text
        if source.endswith('.docx'):
            doc = DocxDocument(source)
            text = "\n".join([p.text for p in doc.paragraphs])
            structure = extract_structure_docx(source)
        elif source.endswith('.pdf'):
            doc = fitz.open(source)
            text = ""
            for page in doc:
                text += page.get_text()
            structure = extract_structure_pdf(source)
        else:
            error("Unsupported file format")
        
        result = find_references(text, structure["sections"])
        print(json.dumps(result, indent=2))
    
    elif args.command == 'resolve-references':
        with open(args.structure, 'r') as f:
            structure = json.load(f)
        with open(args.references, 'r') as f:
            references = json.load(f)
        
        result = resolve_references(structure, references)
        print(json.dumps(result, indent=2))
    
    elif args.command == 'fetch-external':
        result = fetch_external_citations(args.citations)
        print(json.dumps(result, indent=2))
    
    elif args.command == 'build':
        build_html(
            args.source,
            args.output,
            args.resolved,
            args.externals,
            args.title,
            args.compact
        )
    
    return 0


if __name__ == '__main__':
    sys.exit(main())
