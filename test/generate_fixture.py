#!/usr/bin/env python3
"""Generate a synthetic vendor services agreement for testing."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('VENDOR SERVICES AGREEMENT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Parties paragraph
p = doc.add_paragraph()
p.add_run('This Vendor Services Agreement (this "').italic = False
p.add_run('Agreement').bold = True
p.add_run('") is entered into as of January 1, 2024 (the "').italic = False
p.add_run('Effective Date').bold = True
p.add_run('") by and between TechCorp Inc., a Delaware corporation ("').italic = False
p.add_run('Customer').bold = True
p.add_run('"), and ServiceProvider LLC, a California limited liability company ("').italic = False
p.add_run('Vendor').bold = True
p.add_run('").')

# Section 1: Definitions
doc.add_heading('1. Definitions', 1)

doc.add_heading('1.1 Confidential Information', 2)
doc.add_paragraph('"Confidential Information" means all non-public information disclosed by either party to the other, whether in written, oral, electronic, or visual form, including but not limited to technical data, trade secrets, business plans, and customer information.')

doc.add_heading('1.2 Effective Date', 2)
doc.add_paragraph('"Effective Date" has the meaning set forth in the preamble above.')

doc.add_heading('1.3 Services', 2)
doc.add_paragraph('"Services" means the software development and consulting services to be provided by Vendor to Customer as described in Schedule A.')

doc.add_heading('1.4 Permitted Use', 2)
doc.add_paragraph('"Permitted Use" means use of the Confidential Information solely for the purpose of performing obligations or exercising rights under this Agreement.')

doc.add_heading('1.5 Losses', 2)
doc.add_paragraph('"Losses" means any direct damages, costs, and expenses, including reasonable attorneys\' fees.')

doc.add_heading('1.6 Affiliate', 2)
doc.add_paragraph('"Affiliate" means any entity that directly or indirectly controls, is controlled by, or is under common control with a party, where "control" means ownership of more than 50% of voting securities.')

# Section 2: Scope of Services
doc.add_heading('2. Scope of Services', 1)

doc.add_heading('2.1 Services to be Provided', 2)
doc.add_paragraph('Vendor shall provide the Services to Customer in accordance with the specifications set forth in Schedule A. All Services shall be performed in a professional and workmanlike manner.')

doc.add_heading('2.2 Service Levels', 2)
doc.add_paragraph('Vendor shall meet the service level requirements described in Schedule B. Failure to meet such requirements may result in service credits as specified in Section 7.3.')

# Section 3: Fees and Payment
doc.add_heading('3. Fees and Payment', 1)

doc.add_heading('3.1 Fees', 2)
doc.add_paragraph('Customer shall pay Vendor the fees set forth in Schedule C. All fees are non-refundable except as expressly provided in this Agreement.')

doc.add_heading('3.2 Payment Terms', 2)
doc.add_paragraph('All invoices are due within thirty (30) days of receipt. Late payments shall accrue interest at the rate of 1.5% per month, subject to Section 14.6.')

# Section 4: Term and Termination
doc.add_heading('4. Term and Termination', 1)

doc.add_heading('4.1 Term', 2)
doc.add_paragraph('This Agreement shall commence on the Effective Date and continue for an initial term of twelve (12) months, unless earlier terminated as provided herein.')

doc.add_heading('4.2 Termination for Convenience', 2)
doc.add_paragraph('Either party may terminate this Agreement for any reason upon sixty (60) days\' prior written notice to the other party.')

doc.add_heading('4.3 Termination for Cause', 2)
doc.add_paragraph('Either party may terminate this Agreement immediately upon written notice if the other party materially breaches this Agreement and fails to cure such breach within thirty (30) days. The Approved Vendor List shall govern which vendors may be substituted upon termination.')

# Section 5: Confidentiality
doc.add_heading('5. Confidentiality', 1)

doc.add_heading('5.1 Confidentiality Obligations', 2)
doc.add_paragraph('The Receiving Party shall hold all Confidential Information in strict confidence and shall not disclose it to any third party except as permitted under Section 7.3 and Schedule C.')

doc.add_heading('5.2 Permitted Disclosures', 2)
doc.add_paragraph('The Receiving Party may disclose Confidential Information to its employees, contractors, and Affiliates who have a need to know and who are bound by confidentiality obligations no less protective than those set forth herein.')

# Section 6: Intellectual Property
doc.add_heading('6. Intellectual Property', 1)

doc.add_heading('6.1 Ownership', 2)
doc.add_paragraph('All work product created by Vendor in the course of performing the Services shall be the sole and exclusive property of Customer.')

doc.add_heading('6.2 License Grant', 2)
doc.add_paragraph('Vendor hereby grants to Customer a perpetual, irrevocable, worldwide license to use, modify, and distribute any pre-existing Vendor materials incorporated into the work product, for the Permitted Use only.')

# Section 7: Warranties and Disclaimers
doc.add_heading('7. Warranties and Disclaimers', 1)

doc.add_heading('7.1 Performance Warranty', 2)
doc.add_paragraph('Vendor warrants that the Services will be performed in a professional and workmanlike manner in accordance with industry standards.')

doc.add_heading('7.2 Disclaimer', 2)
doc.add_paragraph('EXCEPT AS EXPRESSLY PROVIDED IN SECTION 7.1, VENDOR PROVIDES THE SERVICES "AS IS" WITHOUT WARRANTY OF ANY KIND.')

doc.add_heading('7.3 Remedies', 2)
doc.add_paragraph('Customer\'s sole and exclusive remedy for breach of the warranty in Section 7.1 shall be re-performance of the deficient Services or, if Vendor is unable to cure such deficiency, a pro-rata refund of fees paid for such deficient Services.')

# Section 8: Indemnification
doc.add_heading('8. Indemnification', 1)

doc.add_heading('8.1 Indemnification by Vendor', 2)
doc.add_paragraph('For purposes of this Section 8, "Losses" means direct damages, costs, and reasonable attorneys\' fees arising from third-party claims. Vendor shall indemnify, defend, and hold harmless Customer from any Losses arising from Vendor\'s breach of this Agreement or violation of applicable law, including without limitation 17 U.S.C. Section 512.')

doc.add_heading('8.2 Indemnification Process', 2)
doc.add_paragraph('The indemnified party shall provide prompt written notice of any claim, allow the indemnifying party to control the defense, and provide reasonable cooperation.')

# Section 9: Data Protection
doc.add_heading('9. Data Protection', 1)

doc.add_heading('9.1 Compliance', 2)
doc.add_paragraph('Vendor shall comply with all applicable data protection laws and regulations, including without limitation GDPR Article 28.')

doc.add_heading('9.2 Data Processing', 2)
doc.add_paragraph('To the extent Vendor processes personal data on behalf of Customer, Vendor shall act only on Customer\'s documented instructions and shall implement appropriate technical and organizational measures as required by GDPR Article 32.')

# Section 10: Limitation of Liability
doc.add_heading('10. Limitation of Liability', 1)

doc.add_heading('10.1 Cap on Liability', 2)
doc.add_paragraph('EXCEPT FOR BREACHES OF SECTION 5 (CONFIDENTIALITY) OR SECTION 8 (INDEMNIFICATION), IN NO EVENT SHALL EITHER PARTY\'S TOTAL LIABILITY UNDER THIS AGREEMENT EXCEED THE FEES PAID BY CUSTOMER IN THE TWELVE (12) MONTHS PRECEDING THE CLAIM.')

doc.add_heading('10.2 Exclusion of Damages', 2)
doc.add_paragraph('IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, CONSEQUENTIAL, SPECIAL, OR PUNITIVE DAMAGES, EVEN IF ADVISED OF THE POSSIBILITY THEREOF.')

# Section 11: General Provisions
doc.add_heading('11. General Provisions', 1)

doc.add_heading('11.1 Governing Law', 2)
doc.add_paragraph('This Agreement shall be governed by the laws of the State of California, without regard to its conflicts of law principles.')

doc.add_heading('11.2 Entire Agreement', 2)
doc.add_paragraph('This Agreement, together with all Schedules and Exhibits attached hereto, constitutes the entire agreement between the parties.')

doc.add_heading('11.3 Amendments', 2)
doc.add_paragraph('This Agreement may be amended only by a written instrument signed by both parties.')

doc.add_heading('11.4 Severability', 2)
doc.add_paragraph('If any provision of this Agreement is found to be invalid or unenforceable, the remaining provisions shall continue in full force and effect.')

doc.add_heading('11.5 Notices', 2)
doc.add_paragraph('All notices under this Agreement shall be in writing and delivered to the addresses set forth in Schedule D or such other address as a party may designate by notice.')

# Schedule A
doc.add_page_break()
doc.add_heading('Schedule A - Services Description', 1)
doc.add_paragraph('Vendor shall provide the following services:')
doc.add_paragraph('1. Software development services for the Customer\'s cloud platform')
doc.add_paragraph('2. Technical consulting and architecture review')
doc.add_paragraph('3. Code review and quality assurance')
doc.add_paragraph('4. Documentation and training materials')

# Schedule B
doc.add_heading('Schedule B - Service Level Agreement', 1)
doc.add_paragraph('Response time: 4 hours for critical issues, 24 hours for standard issues')
doc.add_paragraph('Uptime: 99.9% monthly uptime for all delivered services')
doc.add_paragraph('Resolution time: 48 hours for critical bugs, 10 business days for standard bugs')

# Schedule C
doc.add_heading('Schedule C - Fees and Payment Schedule', 1)
doc.add_paragraph('Monthly retainer: $50,000 per month')
doc.add_paragraph('Hourly rate for additional work: $200 per hour')
doc.add_paragraph('Payment terms: Net 30 days from invoice date')

# Save
doc.save('test/fixtures/sample-contract.docx')
print("Generated test/fixtures/sample-contract.docx")
