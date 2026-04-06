#!/usr/bin/env python3
"""Test suite for Xref."""

import json
import os
import sys
import unittest
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent / "skills" / "xref" / "tools"))

import xref


class TestXref(unittest.TestCase):
    
    @classmethod
    def setUpClass(cls):
        """Generate test fixture if it doesn't exist."""
        fixture_path = Path(__file__).parent / "fixtures" / "sample-contract.docx"
        if not fixture_path.exists():
            print("Generating test fixture...")
            import subprocess
            gen_script = Path(__file__).parent / "generate_fixture.py"
            subprocess.run([sys.executable, str(gen_script)], check=True)
        
        cls.fixture_path = str(fixture_path)
    
    def test_extract_structure_docx(self):
        """Test extracting structure from Word document."""
        result = xref.extract_structure_docx(self.fixture_path)
        
        # Check structure
        self.assertIn("sections", result)
        self.assertIn("defined_terms", result)
        self.assertIn("numbering_scheme", result)
        self.assertIn("parser_confidence", result)
        
        # Should have found sections
        self.assertGreater(len(result["sections"]), 0)
        
        # Should have found defined terms
        self.assertGreater(len(result["defined_terms"]), 0)
        
        # Check for specific defined terms
        term_names = [t["term"] for t in result["defined_terms"]]
        self.assertIn("Confidential Information", term_names)
        self.assertIn("Services", term_names)
        
        print(f"\n✓ Extracted {len(result['sections'])} sections and {len(result['defined_terms'])} defined terms")
    
    def test_find_references(self):
        """Test finding references in document."""
        # First extract structure
        structure = xref.extract_structure_docx(self.fixture_path)
        
        # Extract text
        from docx import Document as DocxDocument
        doc = DocxDocument(self.fixture_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Find references
        result = xref.find_references(text, structure["sections"])
        
        # Check structure
        self.assertIn("internal_references", result)
        self.assertIn("defined_term_usages", result)
        self.assertIn("external_citations", result)
        self.assertIn("unresolved", result)
        
        # Should have found some internal references
        self.assertGreater(len(result["internal_references"]), 0)
        
        # Should have found external citations
        self.assertGreater(len(result["external_citations"]), 0)
        
        # Check for specific references
        ref_texts = [r["text"] for r in result["internal_references"]]
        self.assertTrue(any("Section 7.3" in r for r in ref_texts))
        
        # Check for external citations
        ext_texts = [e["text"] for e in result["external_citations"]]
        self.assertTrue(any("GDPR" in e for e in ext_texts))
        
        print(f"\n✓ Found {len(result['internal_references'])} internal refs, "
              f"{len(result['external_citations'])} external citations")
    
    def test_resolve_references(self):
        """Test resolving references to targets."""
        # Extract structure and find references
        structure = xref.extract_structure_docx(self.fixture_path)
        
        from docx import Document as DocxDocument
        doc = DocxDocument(self.fixture_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        references = xref.find_references(text, structure["sections"])
        
        # Resolve
        result = xref.resolve_references(structure, references)
        
        # Check structure
        self.assertIn("internal_references", result)
        self.assertIn("defined_terms", result)
        self.assertIn("unresolved", result)
        
        # Should have resolved some references
        resolved_count = sum(1 for r in result["internal_references"] if r.get("resolved"))
        self.assertGreater(resolved_count, 0)
        
        # Should have some unresolved (broken references in the fixture)
        self.assertGreater(len(result["unresolved"]), 0)
        
        print(f"\n✓ Resolved {resolved_count} references, {len(result['unresolved'])} unresolved")
    
    def test_gdpr_bundled(self):
        """Test that GDPR bundled data loads correctly."""
        from statutes.eu_regulations import get_bundled_gdpr
        
        gdpr = get_bundled_gdpr()
        
        self.assertIn("articles", gdpr)
        self.assertIn("regulation", gdpr)
        
        # Check for specific articles
        self.assertIn("1", gdpr["articles"])
        self.assertIn("28", gdpr["articles"])
        self.assertIn("32", gdpr["articles"])
        
        # Check article structure
        article_28 = gdpr["articles"]["28"]
        self.assertIn("heading", article_28)
        self.assertIn("text", article_28)
        self.assertIn("Processor", article_28["heading"])
        
        print(f"\n✓ GDPR bundled data contains {len(gdpr['articles'])} articles")
    
    def test_cache_operations(self):
        """Test cache get/set operations."""
        from statutes.cache import get_cache_key, get_cached, set_cached
        
        # Create a test cache entry
        cache_key = get_cache_key("test", param1="value1", param2="value2")
        self.assertIsInstance(cache_key, str)
        self.assertEqual(len(cache_key), 64)  # SHA-256 hash
        
        # Test set and get
        test_data = {"test": "data", "value": 123}
        set_cached(cache_key, test_data)
        
        retrieved = get_cached(cache_key)
        self.assertEqual(retrieved, test_data)
        
        # Test non-existent key
        bad_key = get_cache_key("nonexistent", foo="bar")
        self.assertIsNone(get_cached(bad_key))
        
        print("\n✓ Cache operations working")
    
    def test_build_html(self):
        """Test HTML generation."""
        # Generate structure and references
        structure = xref.extract_structure_docx(self.fixture_path)
        
        from docx import Document as DocxDocument
        doc = DocxDocument(self.fixture_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        references = xref.find_references(text, structure["sections"])
        
        # Resolve
        resolved = xref.resolve_references(structure, references)
        
        # Save resolved data to temp file
        import tempfile
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as f:
            resolved_path = f.name
            json.dump(resolved, f)
        
        # Build HTML
        output_path = Path(__file__).parent / "output" / "test-output.html"
        output_path.parent.mkdir(exist_ok=True)
        
        try:
            xref.build_html(
                self.fixture_path,
                str(output_path),
                resolved_path,
                title="Test Contract"
            )
            
            # Check that file was created
            self.assertTrue(output_path.exists())
            
            # Check that it contains expected elements
            with open(output_path, 'r') as f:
                html = f.read()
            
            self.assertIn("Test Contract", html)
            self.assertIn("xrefData", html)
            self.assertIn("Glossary", html)
            self.assertIn("Document Health", html)
            
            print(f"\n✓ Generated HTML output: {output_path}")
            
        finally:
            # Cleanup
            os.unlink(resolved_path)
    
    def test_setup_check(self):
        """Test setup check command."""
        result = xref.setup_check()
        
        # Should return 0 if all deps installed, 1 otherwise
        self.assertIn(result, [0, 1])
        
        print("\n✓ Setup check completed")


if __name__ == '__main__':
    unittest.main(verbosity=2)
